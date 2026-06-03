"""
Generate sample resumes for testing the screening system
Run this file to create test resumes in the resumes folder
"""

from docx import Document
import os
import sys

# Add parent directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import config


def create_resume(name, skills, experience, education):
    """
    Create a single resume document
    
    Parameters:
        name: Candidate's full name
        skills: List of skills
        experience: Experience description
        education: Education description
    """
    doc = Document()
    
    # Header
    doc.add_heading(name, 0)
    
    # Skills section
    doc.add_heading('Technical Skills', level=1)
    for skill in skills:
        doc.add_paragraph(f'• {skill}')
    
    # Experience section
    doc.add_heading('Experience', level=1)
    doc.add_paragraph(experience)
    
    # Education section
    doc.add_heading('Education', level=1)
    doc.add_paragraph(education)
    
    # Save file
    filename = name.lower().replace(' ', '_') + '_resume.docx'
    filepath = os.path.join(config.RESUME_DIR, filename)
    doc.save(filepath)
    print(f"✅ Created: {filename}")


def create_all_resumes():
    """Create multiple sample resumes for testing"""
    
    print("\n" + "="*50)
    print("📝 CREATING SAMPLE RESUMES")
    print("="*50)
    
    candidates = [
        {
            "name": "Alice Johnson",
            "skills": ["Python", "Machine Learning", "NLP", "SQL", "Data Analysis", "Communication"],
            "experience": """Senior Data Scientist at TechCorp (2021-Present)
- Developed ML models for customer prediction using Python and TensorFlow
- Implemented NLP pipelines for text classification
- Led data analysis projects and presented insights to stakeholders""",
            "education": "M.S. in Computer Science, Stanford University (2021)"
        },
        {
            "name": "Bob Smith",
            "skills": ["JavaScript", "HTML", "CSS", "React", "Node.js", "MongoDB"],
            "experience": """Full Stack Developer at WebSolutions (2022-Present)
- Built responsive web applications using React
- Developed REST APIs with Node.js
- Managed MongoDB databases""",
            "education": "B.S. in Information Technology, State University (2022)"
        },
        {
            "name": "Carol Davis",
            "skills": ["Python", "Machine Learning", "NLP", "TensorFlow", "PyTorch", "Deep Learning", "Computer Vision"],
            "experience": """AI Research Scientist at AI Labs (2020-Present)
- Published research papers on NLP and deep learning
- Implemented state-of-the-art models using PyTorch and TensorFlow
- Developed ML pipelines for production systems""",
            "education": "Ph.D. in Artificial Intelligence, MIT (2020)"
        },
        {
            "name": "David Wilson",
            "skills": ["Python", "SQL", "Data Analysis", "Tableau", "Excel", "Statistics", "Communication"],
            "experience": """Data Analyst at DataCorp (2021-Present)
- Analyzed large datasets using Python and SQL
- Created dashboards and reports for management
- Collaborated with cross-functional teams""",
            "education": "B.S. in Statistics, University of Texas (2021)"
        },
        {
            "name": "Emma Brown",
            "skills": ["Java", "C++", "Teamwork", "Problem Solving"],
            "experience": """Junior Developer at Startup (2023-Present)
- Assisted in software development using Java
- Participated in code reviews and team meetings
- Learned agile development methodologies""",
            "education": "B.S. in Computer Engineering, Local University (2023)"
        },
        {
            "name": "John Doe",
            "skills": ["Python", "Machine Learning", "SQL", "Communication", "Data Analysis"],
            "experience": """AI Developer at Tech Company (2022-Present)
- Built machine learning models using Python
- Worked on data analysis projects using SQL
- Collaborated with cross-functional teams""",
            "education": "B.S. in Computer Science, University of Example (2022)"
        }
    ]
    
    # Create each resume
    for candidate in candidates:
        create_resume(
            candidate["name"],
            candidate["skills"],
            candidate["experience"],
            candidate["education"]
        )
    
    print("\n" + "="*50)
    print(f"✅ Created {len(candidates)} sample resumes")
    print(f"📁 Location: {config.RESUME_DIR}")
    print("="*50)
    print("\n🎯 Now run: python main.py")


if __name__ == "__main__":
    create_all_resumes()
